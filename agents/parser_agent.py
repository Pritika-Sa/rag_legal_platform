import os
import re
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Any, Optional
import pdfplumber
from docx import Document
from pydantic import BaseModel, Field
from langchain_core.tools import BaseTool
from langchain_text_splitters import RecursiveCharacterTextSplitter

# Input Schema for LangChain Component
class DocumentParserInput(BaseModel):
    file_path: str = Field(description="The absolute path to the PDF or DOCX file to be parsed")
    document_id: Optional[str] = Field(None, description="Optional unique identifier for the document. If not provided, an MD5 hash of the file will be generated.")
    version: Optional[int] = Field(1, description="Document version number")

# Reusable LangChain Tool Component
class DocumentParsingTool(BaseTool):
    name: str = "document_parser"
    description: str = "Parses PDF and DOCX files, extracts metadata, chunks text using RecursiveCharacterTextSplitter, and returns structured JSON."
    args_schema: type[BaseModel] = DocumentParserInput

    def _run(self, file_path: str, document_id: Optional[str] = None, version: int = 1) -> str:
        """Runs document parsing, chunking, and returns structured JSON output."""
        try:
            result = parse_document_to_json(file_path, document_id, version)
            return json.dumps(result, indent=2)
        except Exception as e:
            return json.dumps({"error": str(e)})

def get_file_md5(file_path: str) -> str:
    """Computes MD5 hash of a file."""
    hasher = hashlib.md5()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def extract_pdf_metadata(file_path: str) -> Dict[str, Any]:
    """Extracts metadata and statistics from PDF file."""
    meta = {}
    try:
        with pdfplumber.open(file_path) as pdf:
            meta["page_count"] = len(pdf.pages)
            # Fetch default metadata properties if present
            if pdf.metadata:
                for k, v in pdf.metadata.items():
                    if isinstance(v, (str, int, float, bool)):
                        meta[k.lower()] = v
    except Exception as e:
        meta["error"] = f"Failed to extract PDF metadata: {str(e)}"
    return meta

def extract_docx_metadata(file_path: str) -> Dict[str, Any]:
    """Extracts metadata and properties from Word document."""
    meta = {}
    try:
        doc = Document(file_path)
        props = doc.core_properties
        meta["author"] = props.author or "Unknown"
        meta["title"] = props.title or "Unknown"
        meta["created"] = str(props.created) if props.created else None
        meta["modified"] = str(props.modified) if props.modified else None
        meta["revision"] = props.revision
    except Exception as e:
        meta["error"] = f"Failed to extract DOCX metadata: {str(e)}"
    return meta

def parse_document_to_json(file_path: str, document_id: Optional[str] = None, version: int = 1) -> Dict[str, Any]:
    """Extracts text and constructs structured JSON chunks and metadata."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    doc_name = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()
    
    # 1. Compute/Assign Document ID and basic info
    if not document_id:
        document_id = get_file_md5(file_path)
        
    upload_date = datetime.now().isoformat()
    
    # 2. Extract text and metadata
    raw_text = ""
    doc_metadata = {
        "file_size_bytes": os.path.getsize(file_path),
        "file_type": ext.replace(".", "")
    }
    
    if ext == ".docx":
        doc = Document(file_path)
        raw_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        doc_metadata.update(extract_docx_metadata(file_path))
    elif ext == ".pdf":
        with pdfplumber.open(file_path) as pdf:
            text_blocks = []
            for page in pdf.pages:
                text_blocks.append(page.extract_text() or "")
            raw_text = "\n".join(text_blocks)
        doc_metadata.update(extract_pdf_metadata(file_path))
    elif ext == ".txt":
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            raw_text = f.read()
    else:
        raise ValueError(f"Unsupported file format: {ext}")
        
    # 3. Create Chunks using RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=150,
        length_function=len
    )
    
    chunks = text_splitter.split_text(raw_text)
    
    # 4. Format structured chunks
    formatted_chunks = []
    for idx, chunk in enumerate(chunks):
        formatted_chunks.append({
            "chunk_id": f"{document_id}_chunk_{idx}",
            "text_content": chunk,
            "metadata": {
                "chunk_index": idx,
                "length": len(chunk)
            }
        })
        
    # 5. Compile structured JSON output representation
    return {
        "document_id": document_id,
        "document_name": doc_name,
        "upload_date": upload_date,
        "version": version,
        "document_metadata": doc_metadata,
        "chunks": formatted_chunks
    }


# Backwards compatibility layer for local orchestrator / regex parsing
def parse_document(file_path: str) -> List[Dict[str, Any]]:
    """Legacy backward-compatible parser segmenting text into sections using regex or fallback chunks."""
    ext = os.path.splitext(file_path)[1].lower()
    section_pattern = re.compile(
        r'^(?:section|clause|article|part)\s+\d+(?:\.\d+)*[:\-\s\.]|^\d+\.\s+[A-Z]', 
        re.IGNORECASE
    )
    
    # Fallback to structured chunking if regex segmentation yields fewer than 3 segments
    try:
        sections = []
        if ext == ".docx":
            doc = Document(file_path)
            current_section = "Preamble"
            current_content = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                if section_pattern.match(text) and len(text) < 150:
                    if current_content:
                        sections.append({
                            "section_name": current_section,
                            "text_content": "\n".join(current_content).strip(),
                            "page_num": None
                        })
                    current_section = text
                    current_content = []
                else:
                    current_content.append(text)
            if current_content:
                sections.append({
                    "section_name": current_section,
                    "text_content": "\n".join(current_content).strip(),
                    "page_num": None
                })
        elif ext == ".pdf":
            current_section = "Preamble"
            current_content = []
            current_page = 1
            with pdfplumber.open(file_path) as pdf:
                for page_idx, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if not text:
                        continue
                    for line in text.split("\n"):
                        line_str = line.strip()
                        if not line_str:
                            continue
                        if section_pattern.match(line_str) and len(line_str) < 150:
                            if current_content:
                                sections.append({
                                    "section_name": current_section,
                                    "text_content": "\n".join(current_content).strip(),
                                    "page_num": current_page
                                })
                            current_section = line_str
                            current_content = []
                            current_page = page_idx + 1
                        else:
                            current_content.append(line_str)
            if current_content:
                sections.append({
                    "section_name": current_section,
                    "text_content": "\n".join(current_content).strip(),
                    "page_num": current_page
                })
        else:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            sections = [{"section_name": "Full Document", "text_content": content, "page_num": 1}]
            
        # If regex didn't find clear sections (e.g. less than 2 sections parsed)
        # fallback to using RecursiveCharacterTextSplitter chunks as sections
        if len(sections) <= 2:
            json_parsed = parse_document_to_json(file_path)
            sections = []
            for ch in json_parsed["chunks"]:
                sections.append({
                    "section_name": f"Section {ch['metadata']['chunk_index'] + 1}",
                    "text_content": ch["text_content"],
                    "page_num": None
                })
        return sections
    except Exception as e:
        # Full fallback to RecursiveCharacterTextSplitter
        print(f"Regex parser failed: {e}. Falling back to RecursiveCharacterTextSplitter...")
        json_parsed = parse_document_to_json(file_path)
        sections = []
        for ch in json_parsed["chunks"]:
            sections.append({
                "section_name": f"Section {ch['metadata']['chunk_index'] + 1}",
                "text_content": ch["text_content"],
                "page_num": None
            })
        return sections
